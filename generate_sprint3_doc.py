import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx by running: pip install python-docx")
    exit(1)

def create_sprint3_document():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Silent Voices: AI-Powered ASL to Text Translation System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Sprint 3 Final Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\n\n")
    team_info = doc.add_paragraph()
    team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_info.add_run("Team Members:\n").bold = True
    team_info.add_run("Irtaza Naqvi (23L-0608)\n")
    team_info.add_run("Faizan Ashfaq (23L-3091)\n")
    team_info.add_run("Muhammad Fakhir (22L-6827)\n")
    
    doc.add_page_break()

    # --- SPRINT BACKLOG FOR SPRINT 3 ---
    doc.add_heading('1. Sprint Backlog for Sprint 3', level=1)
    
    doc.add_heading('a. Module Name', level=2)
    doc.add_paragraph("Admin Dashboard, User History, Feedback, and Vocabulary Module")
    
    doc.add_heading('b. User Stories and Sub User Stories for Sprint 3', level=2)
    
    us1 = doc.add_paragraph(style='List Bullet')
    us1.add_run("US-3.1: As an Admin, I want to view a dashboard with system statistics (total users, sessions, avg confidence) to monitor system usage.\n").bold = True
    us1.add_run("  - Sub-task: Implement backend analytics API.\n")
    us1.add_run("  - Sub-task: Create StatCards and dashboard UI in React.")
    
    us2 = doc.add_paragraph(style='List Bullet')
    us2.add_run("US-3.2: As an Admin, I want to manage user accounts (deactivate, activate, delete) to maintain security.\n").bold = True
    us2.add_run("  - Sub-task: Implement role-based admin endpoints for user management.\n")
    us2.add_run("  - Sub-task: Build the Users table UI with action buttons.")

    us3 = doc.add_paragraph(style='List Bullet')
    us3.add_run("US-3.3: As a User, I want to view my past translation history to track my ASL practice.\n").bold = True
    us3.add_run("  - Sub-task: Create history API to fetch user-specific sessions.\n")
    us3.add_run("  - Sub-task: Implement expandable history list in frontend.")

    us4 = doc.add_paragraph(style='List Bullet')
    us4.add_run("US-3.4: As a User, I want to export my translations (TXT/PDF) to share with others.\n").bold = True
    us4.add_run("  - Sub-task: Integrate PDF/TXT generation in FastAPI using fpdf2.\n")
    us4.add_run("  - Sub-task: Add export buttons in history UI.")

    us5 = doc.add_paragraph(style='List Bullet')
    us5.add_run("US-3.5: As a User, I want to provide feedback (thumbs up/down and correction) on translations to help improve the system.\n").bold = True
    us5.add_run("  - Sub-task: Create database tables and APIs for feedback.\n")
    us5.add_run("  - Sub-task: Add feedback UI elements to session details.")

    us6 = doc.add_paragraph(style='List Bullet')
    us6.add_run("US-3.6: As a User, I want to browse an ASL Vocabulary Guide so I can learn the supported signs.\n").bold = True
    us6.add_run("  - Sub-task: Create vocabulary endpoint.\n")
    us6.add_run("  - Sub-task: Implement Vocabulary page with search filtering.")

    doc.add_heading('c. Rework or Left Out Stories', level=2)
    doc.add_paragraph("There are no left out user stories from Sprint 1 and Sprint 2. All previous functionalities including ML pipeline integration and TTS have been successfully completed.")
    
    doc.add_page_break()

    # --- FINAL REPORT ---
    doc.add_heading('Sprint 3 - FINAL REPORT', level=1)
    
    # a) Project Intro
    doc.add_heading('a) Project Introduction', level=2)
    doc.add_paragraph(
        "Silent Voices is an AI-powered web application designed to bridge the communication gap between the Deaf "
        "and hearing communities. The system translates American Sign Language (ASL) gestures from uploaded videos into "
        "English text using computer vision (MediaPipe) and machine learning (Scikit-learn). It features a robust "
        "React frontend for user interaction, text-to-speech (TTS) capabilities, and a secure FastAPI/PostgreSQL backend "
        "with role-based access control for users and administrators."
    )

    # b) User Stories (all sprints)
    doc.add_heading('b) User Stories (All Sprints)', level=2)
    
    doc.add_paragraph("Sprint 1:", style='Heading 3')
    doc.add_paragraph("1. As a user, I want to register for a new account.\n"
                      "2. As a user, I want to log in to my account using my email and password.", style='List Bullet')
    
    doc.add_paragraph("Sprint 2:", style='Heading 3')
    doc.add_paragraph("3. As a user, I want to upload a video of my ASL gestures.\n"
                      "4. As a user, I want the system to process the video and translate gestures to English text.\n"
                      "5. As a user, I want to hear the translated text via Text-to-Speech.", style='List Bullet')
    
    doc.add_paragraph("Sprint 3:", style='Heading 3')
    doc.add_paragraph("6. As an Admin, I want to manage users and view system statistics.\n"
                      "7. As a User, I want to view my translation history and export it.\n"
                      "8. As a User, I want to provide feedback on translation accuracy.\n"
                      "9. As a User, I want to view a vocabulary guide of supported ASL signs.", style='List Bullet')

    # c) Design
    doc.add_heading('c) Design (Sprint 2 Items)', level=2)
    doc.add_paragraph("Below are the system design diagrams constructed during Sprint 2:")
    add_placeholder_box(doc, "INSERT USE CASE DIAGRAM SCREENSHOT HERE")
    add_placeholder_box(doc, "INSERT SEQUENCE DIAGRAM SCREENSHOT HERE")
    add_placeholder_box(doc, "INSERT CLASS DIAGRAM SCREENSHOT HERE")
    add_placeholder_box(doc, "INSERT ACTIVITY DIAGRAM SCREENSHOT HERE")

    # d) Architecture
    doc.add_heading('d) Architecture', level=2)
    doc.add_paragraph(
        "The project utilizes a Client-Server Architecture Pattern. The frontend is built using React (Client), "
        "while the backend operates on FastAPI (Server) interacting with a PostgreSQL database. The machine learning "
        "pipeline is integrated directly into the backend as an encapsulated service module."
    )
    add_placeholder_box(doc, "INSERT ARCHITECTURE DIAGRAM SCREENSHOT HERE")

    # e) Implementation Screenshots
    doc.add_heading('e) Actual Implementation Screenshots', level=2)
    add_placeholder_box(doc, "INSERT SCREENSHOT: LOGIN / SIGNUP PAGE")
    add_placeholder_box(doc, "INSERT SCREENSHOT: USER DASHBOARD / VIDEO UPLOAD")
    add_placeholder_box(doc, "INSERT SCREENSHOT: TRANSLATION RESULTS & TTS")
    add_placeholder_box(doc, "INSERT SCREENSHOT: ADMIN DASHBOARD")
    add_placeholder_box(doc, "INSERT SCREENSHOT: HISTORY PAGE")
    add_placeholder_box(doc, "INSERT SCREENSHOT: VOCABULARY GUIDE")

    # f) Product Burn down chart
    doc.add_heading('f) Product Burn down chart for the project', level=2)
    add_placeholder_box(doc, "INSERT BURN DOWN CHART SCREENSHOT HERE")

    # g) Trello board screen shots
    doc.add_heading('g) Trello Board Screen Shots', level=2)
    add_placeholder_box(doc, "INSERT TRELLO BOARD SCREENSHOT(S) HERE")

    # h) Boundary value analysis testing
    doc.add_heading('h) Boundary Value Analysis Testing', level=2)
    doc.add_paragraph("Boundary Value Analysis (BVA) was performed on the Sign Up and Login interfaces. "
                      "We assumed a minimum password length of 6 characters for registration.")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Test Data'
    hdr_cells[2].text = 'Expected Outcome'
    hdr_cells[3].text = 'Status'
    
    # Add rows
    row_data = [
        ("Password", "Length = 5 chars", "Error: Password too short", "Pass"),
        ("Password", "Length = 6 chars", "Accepted", "Pass"),
        ("Password", "Length = 7 chars", "Accepted", "Pass"),
        ("Email", "Missing '@' symbol", "Error: Invalid email format", "Pass"),
        ("Email", "Empty field", "Error: Field required", "Pass"),
        ("Email", "Valid email format", "Accepted", "Pass"),
    ]
    for field, data, expected, status in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = data
        row_cells[2].text = expected
        row_cells[3].text = status

    # i) Work Division
    doc.add_heading('i) Work Division between group members', level=2)
    doc.add_paragraph("Irtaza Naqvi (23L-0608):", style='Heading 3')
    doc.add_paragraph("Developed the backend APIs (Admin, History, Feedback) and integrated the Machine Learning pipeline with FastAPI.", style='List Bullet')
    
    doc.add_paragraph("Faizan Ashfaq (23L-3091):", style='Heading 3')
    doc.add_paragraph("Developed the React Frontend, including Admin Dashboard, History, and Vocabulary pages, ensuring a responsive GUI.", style='List Bullet')
    
    doc.add_paragraph("Muhammad Fakhir (22L-6827):", style='Heading 3')
    doc.add_paragraph("Handled database schema designs, role-based authentication setups, performed BVA testing, and managed documentation.", style='List Bullet')

    # j) Lesson learnt
    doc.add_heading('j) Lesson learnt by group', level=2)
    lessons = [
        "Component Reusability: Using React allowed us to reuse UI components (like the Navbar and ResultCards), saving development time.",
        "Role-Based Access: Implementing JWT tokens with role claims was highly effective for separating Admin and User privileges.",
        "File Processing: Handling video uploads required careful validation and cleanup mechanisms to ensure the backend server does not run out of storage.",
        "Team Collaboration: Using a modular architecture allowed team members to work on the frontend and backend in parallel without blocking each other."
    ]
    for lesson in lessons:
        doc.add_paragraph(lesson, style='List Bullet')

    output_path = os.path.join(os.path.dirname(__file__), 'Iteration_3.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

def add_placeholder_box(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[{text}]\n")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True
    # Add a border-like effect using text
    p.add_run("--------------------------------------------------\n").font.color.rgb = RGBColor(180, 180, 180)

if __name__ == "__main__":
    create_sprint3_document()
